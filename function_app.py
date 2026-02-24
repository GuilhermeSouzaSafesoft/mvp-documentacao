import base64
import json
import os
import re
from io import BytesIO
from datetime import datetime

import azure.functions as func
from docx import Document


app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)  # exige Function Key


TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "template.docx")


CODE_REGEX = re.compile(r"^\d{3}\.\d{3}\.\d{3}\.\d{4}$")


def normalize_codigo(codigo_raw: str):
    """
    Aceita:
      - '000.000.000.0000' (já ok)
      - '0000000000000' (13 dígitos) -> formata
      - com espaços/traços -> remove tudo que não é dígito e tenta formatar
    Retorna (codigo_normalizado, aviso_ou_none). Se não der, retorna (None, mensagem_erro).
    """
    if not codigo_raw or not isinstance(codigo_raw, str):
        return None, "Campo 'codigo' inválido."

    codigo_raw = codigo_raw.strip()

    if CODE_REGEX.match(codigo_raw):
        return codigo_raw, None

    digits = re.sub(r"\D", "", codigo_raw)

    if len(digits) == 13:
        formatted = f"{digits[0:3]}.{digits[3:6]}.{digits[6:9]}.{digits[9:13]}"
        return formatted, "O 'codigo' foi normalizado para o formato 000.000.000.0000."
    else:
        return None, "Campo 'codigo' deve estar no formato 000.000.000.0000 (13 dígitos)."


def replace_placeholders_in_doc(doc: Document, mapping: dict):
    """
    Faz replace em:
      - parágrafos
      - tabelas (células)
      - headers/footers

    Observação: placeholders podem quebrar em múltiplos 'runs' no Word.
    Para um MVP robusto, reconstruímos o texto do parágrafo inteiro e regravamos.
    Isso pode perder formatação dentro do parágrafo (negrito em parte do texto).
    Para MVP, costuma ser aceitável. Se você precisar preservar 100% de formatação,
    a solução ideal é usar 'docxtpl' e placeholders estilo Jinja.
    """

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return
        new_text = full_text
        for k, v in mapping.items():
            new_text = new_text.replace(k, v)
        if new_text != full_text:
            # limpa runs e escreve tudo em um run só
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(new_text)

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                for t in cell.tables:
                    replace_in_table(t)

    # corpo
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for t in doc.tables:
        replace_in_table(t)

    # headers/footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for p in header.paragraphs:
            replace_in_paragraph(p)
        for t in header.tables:
            replace_in_table(t)
        for p in footer.paragraphs:
            replace_in_paragraph(p)
        for t in footer.tables:
            replace_in_table(t)


@app.route(route="generate-docx", methods=["POST"])
def generate_docx(req: func.HttpRequest) -> func.HttpResponse:
    try:
        body = req.get_json()
    except ValueError:
        return func.HttpResponse(
            json.dumps({"error": "JSON inválido no corpo da requisição."}, ensure_ascii=False),
            status_code=400,
            mimetype="application/json; charset=utf-8",
        )

    titulo = body.get("titulo")
    codigo = body.get("codigo")

    # validações simples
    missing = []
    if not titulo or not isinstance(titulo, str) or not titulo.strip():
        missing.append("titulo")
    if not codigo or not isinstance(codigo, str) or not codigo.strip():
        missing.append("codigo")
    if missing:
        return func.HttpResponse(
            json.dumps({"error": f"Campos obrigatórios ausentes/invalidos: {', '.join(missing)}."}, ensure_ascii=False),
            status_code=400,
            mimetype="application/json; charset=utf-8",
        )

    titulo = titulo.strip()
    codigo_norm, msg = normalize_codigo(codigo.strip())
    if codigo_norm is None:
        return func.HttpResponse(
            json.dumps({"error": msg}, ensure_ascii=False),
            status_code=400,
            mimetype="application/json; charset=utf-8",
        )

    if not os.path.exists(TEMPLATE_PATH):
        return func.HttpResponse(
            json.dumps({"error": "Template .docx não encontrado no deploy (templates/template.docx)."}, ensure_ascii=False),
            status_code=500,
            mimetype="application/json; charset=utf-8",
        )

    # carrega template e substitui placeholders
    doc = Document(TEMPLATE_PATH)
    mapping = {
        "{{TITULO}}": titulo,
        "{{CODIGO}}": codigo_norm,
    }
    replace_placeholders_in_doc(doc, mapping)

    # salva em memória
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    file_bytes = out.read()

    # nome do arquivo (seguro e simples)
    safe_title = re.sub(r"[^a-zA-Z0-9_\- ]", "", titulo)[:50].strip().replace(" ", "_")
    timestamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    filename = f"{codigo_norm}_{safe_title}_{timestamp}.docx" if safe_title else f"{codigo_norm}_{timestamp}.docx"

    headers = {
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "Content-Disposition": f'attachment; filename="{filename}"'
    }
    # Se quiser avisar que normalizou o código, você pode mandar um header custom:
    if msg:
        headers["X-Warning"] = msg

    return func.HttpResponse(
        body=file_bytes,
        status_code=200,
        headers=headers
    )
